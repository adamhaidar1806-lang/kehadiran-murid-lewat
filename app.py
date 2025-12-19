import os
import io
from datetime import datetime, date, timedelta
from functools import wraps

from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file, Response
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash
from sqlalchemy import func, or_, and_, extract
from sqlalchemy.orm import joinedload
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

from models import db, User, Tingkatan, Kelas, Murid, CategoryAlasan, KehadiranLewat, Denda, ActivityLog, SuratAmaran

app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET") or os.environ.get("FLASK_SECRET_KEY") or "sistem-kehadiran-lewat-secret-key-2024"
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("DATABASE_URL")
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(hours=8)

db.init_app(app)

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'index'
login_manager.login_message = 'Sila log masuk untuk mengakses halaman ini.'
login_manager.login_message_category = 'warning'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role != 'admin':
            flash('Anda tidak mempunyai akses ke halaman ini.', 'danger')
            return redirect(url_for('dashboard_overview'))
        return f(*args, **kwargs)
    return decorated_function

def log_activity(action, description=None):
    if current_user.is_authenticated:
        log = ActivityLog(
            user_id=current_user.id,
            action=action,
            description=description
        )
        db.session.add(log)
        db.session.commit()

def get_category_by_reason(reason):
    if not reason:
        return None

    reason_lower = reason.lower()

    cuaca_keywords = ['hujan', 'panas', 'ribut', 'banjir', 'kilat', 'petir', 'sejuk', 'lebat', 'cuaca']
    keluarga_keywords = ['sakit', 'hospital', 'kecemasan', 'emergency', 'ibu', 'bapa', 'adik', 'kakak', 
                         'nenek', 'datuk', 'hantar', 'penghantar', 'keluarga', 'rumah']
    transport_keywords = ['jalan', 'sesak', 'jam', 'traffic', 'kereta', 'motor', 'motosikal', 'bas', 
                          'rosak', 'pancit', 'tayar', 'minyak', 'kemalangan', 'accident', 'kenderaan']

    category = CategoryAlasan.query.filter_by(nama='Lain-lain').first()

    if any(k in reason_lower for k in cuaca_keywords):
        category = CategoryAlasan.query.filter_by(nama='Iklim/Cuaca').first()
    elif any(k in reason_lower for k in keluarga_keywords):
        category = CategoryAlasan.query.filter_by(nama='Masalah Keluarga').first()
    elif any(k in reason_lower for k in transport_keywords):
        category = CategoryAlasan.query.filter_by(nama='Masalah Transport').first()

    return category

def get_weekly_stats():
    today = date.today()
    start_of_week = today - timedelta(days=today.weekday())

    total = KehadiranLewat.query.filter(
        KehadiranLewat.tarikh >= start_of_week,
        KehadiranLewat.tarikh <= today
    ).count()

    lelaki = db.session.query(KehadiranLewat).join(Murid).filter(
        KehadiranLewat.tarikh >= start_of_week,
        KehadiranLewat.tarikh <= today,
        Murid.jantina == 'Lelaki'
    ).count()

    perempuan = db.session.query(KehadiranLewat).join(Murid).filter(
        KehadiranLewat.tarikh >= start_of_week,
        KehadiranLewat.tarikh <= today,
        Murid.jantina == 'Perempuan'
    ).count()

    return {'total': total, 'lelaki': lelaki, 'perempuan': perempuan}

def get_murid_with_warnings(month=None, year=None):
    if month is None:
        month = date.today().month
    if year is None:
        year = date.today().year

    subquery = db.session.query(
        KehadiranLewat.murid_id,
        func.count(KehadiranLewat.id).label('count')
    ).filter(
        extract('month', KehadiranLewat.tarikh) == month,
        extract('year', KehadiranLewat.tarikh) == year
    ).group_by(KehadiranLewat.murid_id).having(func.count(KehadiranLewat.id) >= 3).subquery()

    murid_list = db.session.query(Murid, subquery.c.count).join(
        subquery, Murid.id == subquery.c.murid_id
    ).options(joinedload(Murid.kelas)).all()

    result = []
    for murid, count in murid_list:
        kehadiran = KehadiranLewat.query.filter(
            KehadiranLewat.murid_id == murid.id,
            extract('month', KehadiranLewat.tarikh) == month,
            extract('year', KehadiranLewat.tarikh) == year
        ).order_by(KehadiranLewat.tarikh).all()

        result.append({
            'murid': murid,
            'count': count,
            'kehadiran': kehadiran
        })

    return result

def init_database():
    db.create_all()

    if User.query.count() == 0:
        admin = User(
            username='admin',
            role='admin',
            nama_guru='Pengetua'
        )
        admin.set_password('skuses7620')
        db.session.add(admin)

        guru = User(
            username='guru',
            role='guru',
            nama_guru='Guru Bertugas'
        )
        guru.set_password('smkserikundang7620')
        db.session.add(guru)

        db.session.commit()

    if Tingkatan.query.count() == 0:
        for i in range(1, 6):
            tingkatan = Tingkatan(nama=f'Tingkatan {i}')
            db.session.add(tingkatan)
        db.session.commit()

    if CategoryAlasan.query.count() == 0:
        categories = [
            CategoryAlasan(nama='Iklim/Cuaca', keywords='hujan,panas,ribut,banjir,kilat,petir,sejuk'),
            CategoryAlasan(nama='Masalah Keluarga', keywords='sakit,hospital,ibu,bapa,keluarga,emergency'),
            CategoryAlasan(nama='Masalah Transport', keywords='jalan,sesak,kereta,motor,pancit,rosak,traffic'),
            CategoryAlasan(nama='Lain-lain', keywords='')
        ]
        for cat in categories:
            db.session.add(cat)
        db.session.commit()

with app.app_context():
    init_database()

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard_overview'))
    stats = get_weekly_stats()
    return render_template('index.html', stats=stats)

@app.route('/login', methods=['POST'])
def login():
    username = request.form.get('username', '').strip()
    password = request.form.get('password', '')

    user = User.query.filter_by(username=username).first()

    if user and user.check_password(password):
        login_user(user, remember=True)
        log_activity('login', f'Pengguna {username} log masuk')
        flash(f'Selamat datang, {user.nama_guru or user.username}!', 'success')
        return redirect(url_for('dashboard_overview'))

    flash('Nama pengguna atau kata laluan tidak sah.', 'danger')
    return redirect(url_for('index'))

@app.route('/logout')
@login_required
def logout():
    log_activity('logout', f'Pengguna {current_user.username} log keluar')
    logout_user()
    flash('Anda telah berjaya log keluar.', 'info')
    return redirect(url_for('index'))

@app.route('/dashboard')
@login_required
def dashboard_overview():
    stats = get_weekly_stats()

    today = date.today()
    month = today.month
    year = today.year

    monthly_stats = db.session.query(
        extract('day', KehadiranLewat.tarikh).label('day'),
        func.count(KehadiranLewat.id).label('count')
    ).filter(
        extract('month', KehadiranLewat.tarikh) == month,
        extract('year', KehadiranLewat.tarikh) == year
    ).group_by(extract('day', KehadiranLewat.tarikh)).all()

    category_stats = db.session.query(
        CategoryAlasan.nama,
        func.count(KehadiranLewat.id).label('count')
    ).join(KehadiranLewat, CategoryAlasan.id == KehadiranLewat.category_id).filter(
        extract('month', KehadiranLewat.tarikh) == month,
        extract('year', KehadiranLewat.tarikh) == year
    ).group_by(CategoryAlasan.nama).all()

    top_kelas = db.session.query(
        Kelas.nama_kelas,
        func.count(KehadiranLewat.id).label('count')
    ).join(Murid, Kelas.id == Murid.kelas_id).join(
        KehadiranLewat, Murid.id == KehadiranLewat.murid_id
    ).filter(
        extract('month', KehadiranLewat.tarikh) == month,
        extract('year', KehadiranLewat.tarikh) == year
    ).group_by(Kelas.nama_kelas).order_by(func.count(KehadiranLewat.id).desc()).limit(5).all()

    warnings_count = len(get_murid_with_warnings())

    recent_checkins = KehadiranLewat.query.options(
        joinedload(KehadiranLewat.murid)
    ).order_by(KehadiranLewat.created_at.desc()).limit(10).all()

    return render_template('dashboard_overview.html',
                         stats=stats,
                         monthly_stats=monthly_stats,
                         category_stats=category_stats,
                         top_kelas=top_kelas,
                         warnings_count=warnings_count,
                         recent_checkins=recent_checkins,
                         current_month=today.strftime('%B %Y'))

@app.route('/dashboard/amaran')
@login_required
def dashboard_amaran():
    month = request.args.get('month', date.today().month, type=int)
    year = request.args.get('year', date.today().year, type=int)

    warnings = get_murid_with_warnings(month, year)

    return render_template('dashboard_amaran.html',
                         warnings=warnings,
                         month=month,
                         year=year)

@app.route('/dashboard/checkin', methods=['GET', 'POST'])
@login_required
def dashboard_checkin():
    if request.method == 'POST':
        murid_id = request.form.get('murid_id')
        nama_penuh = request.form.get('nama_penuh', '').strip()
        ic = request.form.get('ic', '').strip()
        jantina = request.form.get('jantina', '').strip()
        kelas_id = request.form.get('kelas_id')
        alasan = request.form.get('alasan', '').strip()
        nota = request.form.get('nota', '').strip()
        minit_lewat = request.form.get('minit_lewat', 0, type=int)

        now = datetime.now()

        if murid_id:
            murid = Murid.query.get(murid_id)
        else:
            existing = Murid.query.filter(
                func.lower(Murid.ic) == ic.lower()
            ).first()

            if existing:
                murid = existing
            else:
                murid = Murid(
                    nama_penuh=nama_penuh,
                    ic=ic,
                    jantina=jantina,
                    kelas_id=kelas_id
                )
                db.session.add(murid)
                db.session.flush()

        category = get_category_by_reason(alasan)

        kehadiran = KehadiranLewat(
            murid_id=murid.id,
            tarikh=now.date(),
            masa_sampai=now.time(),
            minit_lewat=minit_lewat,
            alasan=alasan,
            category_id=category.id if category else None,
            nota=nota,
            checked_by=current_user.id
        )
        db.session.add(kehadiran)
        db.session.commit()

        log_activity('checkin', f'Check-in murid: {murid.nama_penuh}')

        count_this_month = KehadiranLewat.query.filter(
            KehadiranLewat.murid_id == murid.id,
            extract('month', KehadiranLewat.tarikh) == now.month,
            extract('year', KehadiranLewat.tarikh) == now.year
        ).count()

        if count_this_month == 3:
            flash(f'AMARAN: {murid.nama_penuh} telah lewat 3 kali bulan ini!', 'warning')
        else:
            flash(f'Check-in berjaya untuk {murid.nama_penuh}. (Lewat kali ke-{count_this_month} bulan ini)', 'success')

        return redirect(url_for('dashboard_checkin'))

    tingkatan_list = Tingkatan.query.all()
    kelas_list = Kelas.query.all()
    categories = CategoryAlasan.query.all()

    return render_template('dashboard_checkin.html',
                         tingkatan_list=tingkatan_list,
                         kelas_list=kelas_list,
                         categories=categories)

@app.route('/dashboard/history')
@login_required
def dashboard_history():
    filter_type = request.args.get('filter', 'weekly')
    filter_date = request.args.get('date')
    filter_month = request.args.get('month', type=int)
    filter_year = request.args.get('year', type=int)
    filter_kelas = request.args.get('kelas')
    filter_jantina = request.args.get('jantina')
    filter_nama = request.args.get('nama', '').strip()

    today = date.today()
    query = KehadiranLewat.query.options(
        joinedload(KehadiranLewat.murid).joinedload(Murid.kelas),
        joinedload(KehadiranLewat.category)
    )

    if filter_type == 'weekly':
        start_of_week = today - timedelta(days=today.weekday())
        query = query.filter(KehadiranLewat.tarikh >= start_of_week)
        date_range = f"{start_of_week.strftime('%d/%m/%Y')} - {today.strftime('%d/%m/%Y')}"
    elif filter_type == 'monthly':
        if filter_month and filter_year:
            query = query.filter(
                extract('month', KehadiranLewat.tarikh) == filter_month,
                extract('year', KehadiranLewat.tarikh) == filter_year
            )
            date_range = f"{filter_month}/{filter_year}"
        else:
            query = query.filter(
                extract('month', KehadiranLewat.tarikh) == today.month,
                extract('year', KehadiranLewat.tarikh) == today.year
            )
            date_range = today.strftime('%B %Y')
    elif filter_type == 'date' and filter_date:
        try:
            specific_date = datetime.strptime(filter_date, '%Y-%m-%d').date()
            query = query.filter(KehadiranLewat.tarikh == specific_date)
            date_range = specific_date.strftime('%d/%m/%Y')
        except:
            date_range = "Tarikh tidak sah"
    else:
        date_range = "Semua rekod"

    if filter_kelas:
        query = query.join(Murid).join(Kelas).filter(Kelas.nama_kelas == filter_kelas)

    if filter_jantina:
        if not filter_kelas:
            query = query.join(Murid)
        query = query.filter(Murid.jantina == filter_jantina)

    if filter_nama:
        if not filter_kelas and not filter_jantina:
            query = query.join(Murid)
        query = query.filter(func.lower(Murid.nama_penuh).contains(filter_nama.lower()))

    records = query.order_by(KehadiranLewat.tarikh.desc(), KehadiranLewat.masa_sampai.desc()).all()

    total = len(records)
    lelaki = sum(1 for r in records if r.murid.jantina == 'Lelaki')
    perempuan = sum(1 for r in records if r.murid.jantina == 'Perempuan')

    kelas_list = Kelas.query.all()

    return render_template('dashboard_history.html',
                         records=records,
                         filter_type=filter_type,
                         date_range=date_range,
                         total=total,
                         lelaki=lelaki,
                         perempuan=perempuan,
                         kelas_list=kelas_list,
                         filter_kelas=filter_kelas,
                         filter_jantina=filter_jantina,
                         filter_nama=filter_nama)

@app.route('/dashboard/murid')
@login_required
@admin_required
def dashboard_murid():
    tingkatan_list = Tingkatan.query.options(
        joinedload(Tingkatan.kelas)
    ).all()

    return render_template('dashboard_murid.html', tingkatan_list=tingkatan_list)

@app.route('/dashboard/murid/tingkatan/<int:tingkatan_id>')
@login_required
@admin_required
def view_tingkatan(tingkatan_id):
    tingkatan = Tingkatan.query.get_or_404(tingkatan_id)
    kelas_list = Kelas.query.filter_by(tingkatan_id=tingkatan_id).all()

    return render_template('dashboard_tingkatan.html',
                         tingkatan=tingkatan,
                         kelas_list=kelas_list)

@app.route('/dashboard/murid/kelas/<int:kelas_id>')
@login_required
@admin_required
def view_kelas(kelas_id):
    kelas = Kelas.query.get_or_404(kelas_id)
    murid_list = Murid.query.filter_by(kelas_id=kelas_id, is_deleted=False).all()
    all_kelas = Kelas.query.filter(Kelas.id != kelas_id).all()

    return render_template('dashboard_kelas.html',
                         kelas=kelas,
                         murid_list=murid_list,
                         all_kelas=all_kelas)

@app.route('/dashboard/murid/kelas/add', methods=['POST'])
@login_required
@admin_required
def add_kelas():
    nama_kelas = request.form.get('nama_kelas', '').strip()
    nama_guru_kelas = request.form.get('nama_guru_kelas', '').strip()
    tingkatan_id = request.form.get('tingkatan_id', type=int)

    if not nama_kelas or not tingkatan_id:
        flash('Sila isi semua maklumat yang diperlukan.', 'danger')
        return redirect(request.referrer or url_for('dashboard_murid'))

    kelas = Kelas(
        nama_kelas=nama_kelas,
        nama_guru_kelas=nama_guru_kelas,
        tingkatan_id=tingkatan_id
    )
    db.session.add(kelas)
    db.session.commit()

    log_activity('add_kelas', f'Tambah kelas: {nama_kelas}')
    flash(f'Kelas {nama_kelas} berjaya ditambah.', 'success')
    return redirect(url_for('view_tingkatan', tingkatan_id=tingkatan_id))

@app.route('/dashboard/murid/kelas/<int:kelas_id>/delete')
@login_required
@admin_required
def delete_kelas(kelas_id):
    kelas = Kelas.query.get_or_404(kelas_id)
    tingkatan_id = kelas.tingkatan_id
    nama = kelas.nama_kelas

    db.session.delete(kelas)
    db.session.commit()

    log_activity('delete_kelas', f'Padam kelas: {nama}')
    flash(f'Kelas {nama} berjaya dipadam.', 'success')
    return redirect(url_for('view_tingkatan', tingkatan_id=tingkatan_id))

@app.route('/dashboard/murid/add', methods=['POST'])
@login_required
@admin_required
def add_murid():
    nama_penuh = request.form.get('nama_penuh', '').strip()
    ic = request.form.get('ic', '').strip()
    jantina = request.form.get('jantina', '').strip()
    no_ibu_bapa = request.form.get('no_ibu_bapa', '').strip()
    kelas_id = request.form.get('kelas_id', type=int)

    if not nama_penuh or not ic or not jantina or not kelas_id:
        flash('Sila isi semua maklumat yang diperlukan.', 'danger')
        return redirect(request.referrer or url_for('dashboard_murid'))

    existing = Murid.query.filter(func.lower(Murid.ic) == ic.lower()).first()
    if existing:
        flash(f'Murid dengan IC {ic} sudah wujud dalam sistem.', 'warning')
        return redirect(request.referrer or url_for('dashboard_murid'))

    murid = Murid(
        nama_penuh=nama_penuh,
        ic=ic,
        jantina=jantina,
        no_ibu_bapa=no_ibu_bapa,
        kelas_id=kelas_id
    )
    db.session.add(murid)
    db.session.commit()

    log_activity('add_murid', f'Tambah murid: {nama_penuh}')
    flash(f'Murid {nama_penuh} berjaya ditambah.', 'success')
    return redirect(url_for('view_kelas', kelas_id=kelas_id))

@app.route('/dashboard/murid/<int:murid_id>/edit', methods=['POST'])
@login_required
@admin_required
def edit_murid(murid_id):
    murid = Murid.query.get_or_404(murid_id)

    murid.nama_penuh = request.form.get('nama_penuh', murid.nama_penuh).strip()
    murid.ic = request.form.get('ic', murid.ic).strip()
    murid.jantina = request.form.get('jantina', murid.jantina).strip()
    murid.no_ibu_bapa = request.form.get('no_ibu_bapa', '').strip()

    db.session.commit()

    log_activity('edit_murid', f'Edit murid: {murid.nama_penuh}')
    flash(f'Maklumat {murid.nama_penuh} berjaya dikemaskini.', 'success')
    return redirect(url_for('view_kelas', kelas_id=murid.kelas_id))

@app.route('/dashboard/murid/<int:murid_id>/delete')
@login_required
@admin_required
def delete_murid(murid_id):
    murid = Murid.query.get_or_404(murid_id)
    kelas_id = murid.kelas_id
    nama = murid.nama_penuh

    murid.is_deleted = True
    db.session.commit()

    log_activity('delete_murid', f'Padam murid: {nama}')
    flash(f'Murid {nama} berjaya dipadam.', 'success')
    return redirect(url_for('view_kelas', kelas_id=kelas_id))

@app.route('/dashboard/murid/<int:murid_id>/pindah', methods=['POST'])
@login_required
@admin_required
def pindah_murid(murid_id):
    murid = Murid.query.get_or_404(murid_id)
    action = request.form.get('action')

    if action == 'pindah_sekolah':
        nama = murid.nama_penuh
        murid.is_deleted = True
        db.session.commit()
        log_activity('pindah_sekolah', f'Murid pindah sekolah: {nama}')
        flash(f'Murid {nama} telah dipindahkan ke sekolah lain.', 'success')
    elif action == 'pindah_kelas':
        new_kelas_id = request.form.get('new_kelas_id', type=int)
        if new_kelas_id:
            old_kelas = murid.kelas.nama_kelas
            murid.kelas_id = new_kelas_id
            db.session.commit()
            new_kelas = Kelas.query.get(new_kelas_id)
            log_activity('pindah_kelas', f'Murid {murid.nama_penuh} pindah dari {old_kelas} ke {new_kelas.nama_kelas}')
            flash(f'{murid.nama_penuh} berjaya dipindahkan ke {new_kelas.nama_kelas}.', 'success')

    return redirect(request.referrer or url_for('dashboard_murid'))

@app.route('/dashboard/profile', methods=['GET', 'POST'])
@login_required
def dashboard_profile():
    if request.method == 'POST':
        action = request.form.get('action')

        if action == 'update_profile':
            current_user.nama_guru = request.form.get('nama_guru', '').strip()
            bertugas_dari = request.form.get('bertugas_dari')
            bertugas_hingga = request.form.get('bertugas_hingga')

            if bertugas_dari:
                current_user.bertugas_dari = datetime.strptime(bertugas_dari, '%Y-%m-%d').date()
            if bertugas_hingga:
                current_user.bertugas_hingga = datetime.strptime(bertugas_hingga, '%Y-%m-%d').date()

            db.session.commit()
            log_activity('update_profile', 'Kemaskini profil')
            flash('Profil berjaya dikemaskini.', 'success')

        elif action == 'change_password':
            current_password = request.form.get('current_password')
            new_password = request.form.get('new_password')
            confirm_password = request.form.get('confirm_password')

            if not current_user.check_password(current_password):
                flash('Kata laluan semasa tidak betul.', 'danger')
            elif new_password != confirm_password:
                flash('Kata laluan baru tidak sepadan.', 'danger')
            elif len(new_password) < 6:
                flash('Kata laluan baru mesti sekurang-kurangnya 6 aksara.', 'danger')
            else:
                current_user.set_password(new_password)
                db.session.commit()
                log_activity('change_password', 'Tukar kata laluan')
                flash('Kata laluan berjaya ditukar.', 'success')

        return redirect(url_for('dashboard_profile'))

    return render_template('dashboard_profile.html')

@app.route('/api/search-murid')
@login_required
def api_search_murid():
    query = request.args.get('q', '').strip()

    if len(query) < 2:
        return jsonify([])

    murid_list = Murid.query.join(Kelas).filter(
        Murid.is_deleted == False,
        or_(
            func.lower(Murid.nama_penuh).contains(query.lower()),
            func.lower(Murid.ic).contains(query.lower())
        )
    ).limit(10).all()

    result = []
    for murid in murid_list:
        result.append({
            'id': murid.id,
            'nama_penuh': murid.nama_penuh,
            'ic': murid.ic,
            'kelas': murid.kelas.nama_kelas,
            'jantina': murid.jantina
        })

    return jsonify(result)

@app.route('/generate-surat/<int:murid_id>')
@login_required
def generate_surat(murid_id):
    murid = Murid.query.get_or_404(murid_id)
    month = request.args.get('month', date.today().month, type=int)
    year = request.args.get('year', date.today().year, type=int)
    preview = request.args.get('preview', False)

    kehadiran = KehadiranLewat.query.filter(
        KehadiranLewat.murid_id == murid_id,
        extract('month', KehadiranLewat.tarikh) == month,
        extract('year', KehadiranLewat.tarikh) == year
    ).order_by(KehadiranLewat.tarikh).all()

    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run(f"Tarikh: {date.today().strftime('%d/%m/%Y')}")

    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Perkara: Surat Amaran – Kehadiran Lewat Melebihi 3 Kali")
    title_run.bold = True

    doc.add_paragraph()

    doc.add_paragraph("Kepada:")
    doc.add_paragraph(f"Nama Penuh Murid: {murid.nama_penuh}")
    doc.add_paragraph(f"Kelas: {murid.kelas.nama_kelas}")

    doc.add_paragraph()

    doc.add_paragraph("Assalamualaikum dan Salam sejahtera")

    doc.add_paragraph()

    bulan_names = ['', 'Januari', 'Februari', 'Mac', 'April', 'Mei', 'Jun', 
                   'Julai', 'Ogos', 'September', 'Oktober', 'November', 'Disember']
    bulan_name = bulan_names[month]

    content = doc.add_paragraph()
    content.add_run(
        f"Dengan segala hormatnya, saya selaku guru bertugas mingguan ingin memaklumkan bahawa "
        f"murid yang dinyatakan di atas telah hadir lewat ke sekolah melebihi tiga (3) kali "
        f"bagi bulan {bulan_name} {year}. Butiran kelewatan adalah seperti berikut:"
    )

    doc.add_paragraph()

    hari_names = ['Isnin', 'Selasa', 'Rabu', 'Khamis', 'Jumaat', 'Sabtu', 'Ahad']

    for k in kehadiran:
        hari = hari_names[k.tarikh.weekday()]
        tarikh = k.tarikh.strftime('%d/%m/%Y')
        masa = k.masa_sampai.strftime('%H:%M')
        doc.add_paragraph(f"• {hari}, {tarikh}, {masa}")

    doc.add_paragraph()

    doc.add_paragraph(
        "Kelewatan berulang ini adalah dikesan melalui guru bertugas mingguan dan dicatat "
        "sebagai tindakan disiplin yang perlu diberi perhatian. Murid diminta untuk mengambil "
        "perhatian serius terhadap perkara ini dan memastikan menghadirkan diri tepat pada "
        "waktunya pada masa hadapan bagi mengelakkan tindakan disiplin seterusnya."
    )

    doc.add_paragraph()

    doc.add_paragraph("Sekian, terima kasih atas kerjasama pihak murid dan ibubapa/penjaga.")

    doc.add_paragraph()
    doc.add_paragraph()

    signature = doc.add_paragraph()
    signature.add_run("Yang benar,")
    doc.add_paragraph()
    doc.add_paragraph()

    guru_name = current_user.nama_guru or current_user.username
    guru_para = doc.add_paragraph()
    guru_run = guru_para.add_run(f"({guru_name})")
    guru_run.bold = True
    doc.add_paragraph("Guru Bertugas Mingguan")

    surat_record = SuratAmaran(
        murid_id=murid_id,
        bulan=month,
        tahun=year,
        printed_by=current_user.id
    )
    db.session.add(surat_record)
    db.session.commit()

    log_activity('print_surat', f'Print surat amaran untuk {murid.nama_penuh}')

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"Surat_Amaran_{murid.nama_penuh.replace(' ', '_')}_{month}_{year}.docx"

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/dashboard/denda/<int:murid_id>', methods=['POST'])
@login_required
def add_denda(murid_id):
    murid = Murid.query.get_or_404(murid_id)

    jenis_denda = request.form.get('jenis_denda', '').strip()
    nota = request.form.get('nota', '').strip()

    if not jenis_denda:
        flash('Sila masukkan jenis denda.', 'danger')
        return redirect(request.referrer)

    denda = Denda(
        murid_id=murid_id,
        jenis_denda=jenis_denda,
        tarikh=date.today(),
        nota=nota,
        assigned_by=current_user.id
    )
    db.session.add(denda)
    db.session.commit()

    log_activity('add_denda', f'Tambah denda untuk {murid.nama_penuh}: {jenis_denda}')
    flash(f'Denda berjaya ditambah untuk {murid.nama_penuh}.', 'success')

    return redirect(request.referrer or url_for('dashboard_amaran'))

@app.route('/export/csv')
@login_required
def export_csv():
    filter_type = request.args.get('filter', 'weekly')
    filter_month = request.args.get('month', type=int)
    filter_year = request.args.get('year', type=int)
    filter_date = request.args.get('date')
    filter_kelas = request.args.get('kelas', '').strip()
    filter_jantina = request.args.get('jantina', '').strip()
    filter_nama = request.args.get('nama', '').strip()

    today = date.today()
    query = KehadiranLewat.query.options(
        joinedload(KehadiranLewat.murid).joinedload(Murid.kelas),
        joinedload(KehadiranLewat.category)
    )

    if filter_type == 'weekly':
        start_of_week = today - timedelta(days=today.weekday())
        query = query.filter(KehadiranLewat.tarikh >= start_of_week)
        title = f"Mingguan_{start_of_week.strftime('%d%m%Y')}_{today.strftime('%d%m%Y')}"
    elif filter_type == 'monthly' and filter_month and filter_year:
        query = query.filter(
            extract('month', KehadiranLewat.tarikh) == filter_month,
            extract('year', KehadiranLewat.tarikh) == filter_year
        )
        title = f"Bulanan_{filter_month}_{filter_year}"
    elif filter_type == 'date' and filter_date:
        try:
            specific_date = datetime.strptime(filter_date, '%Y-%m-%d').date()
            query = query.filter(KehadiranLewat.tarikh == specific_date)
            title = f"Tarikh_{filter_date}"
        except:
            title = "Semua_Rekod"
    else:
        title = "Semua_Rekod"

    if filter_kelas:
        query = query.join(Murid).join(Kelas).filter(Kelas.nama_kelas == filter_kelas)

    if filter_jantina:
        if not filter_kelas:
            query = query.join(Murid)
        query = query.filter(Murid.jantina == filter_jantina)

    if filter_nama:
        if not filter_kelas and not filter_jantina:
            query = query.join(Murid)
        query = query.filter(func.lower(Murid.nama_penuh).contains(filter_nama.lower()))

    records = query.order_by(KehadiranLewat.tarikh.desc()).all()

    data = []
    for r in records:
        data.append({
            'Tarikh': r.tarikh.strftime('%d/%m/%Y'),
            'Masa Sampai': r.masa_sampai.strftime('%H:%M'),
            'Minit Lewat': r.minit_lewat or 0,
            'Nama Murid': r.murid.nama_penuh,
            'IC': r.murid.ic,
            'Kelas': r.murid.kelas.nama_kelas,
            'Jantina': r.murid.jantina,
            'Kategori Alasan': r.category.nama if r.category else '',
            'Alasan': r.alasan or '',
            'Nota': r.nota or ''
        })

    df = pd.DataFrame(data)

    total = len(data)
    lelaki = sum(1 for d in data if d['Jantina'] == 'Lelaki')
    perempuan = sum(1 for d in data if d['Jantina'] == 'Perempuan')

    csv_output = io.StringIO()
    df.to_csv(csv_output, index=False, encoding='utf-8')
    csv_output.seek(0)

    return Response(
        csv_output.getvalue(),
        mimetype='text/csv; charset=utf-8',
        headers={'Content-Disposition': f'attachment; filename="Laporan_Kehadiran_Lewat_{title}.csv"'}
    )

@app.route('/export/pdf')
@login_required
def export_pdf():
    filter_type = request.args.get('filter', 'monthly')
    filter_month = request.args.get('month', date.today().month, type=int)
    filter_year = request.args.get('year', date.today().year, type=int)

    query = KehadiranLewat.query.options(
        joinedload(KehadiranLewat.murid).joinedload(Murid.kelas),
        joinedload(KehadiranLewat.category)
    ).filter(
        extract('month', KehadiranLewat.tarikh) == filter_month,
        extract('year', KehadiranLewat.tarikh) == filter_year
    )

    records = query.order_by(KehadiranLewat.tarikh).all()

    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    p.setFont("Helvetica-Bold", 16)
    p.drawCentredString(width/2, height - 50, "Laporan Kehadiran Lewat")

    bulan_names = ['', 'Januari', 'Februari', 'Mac', 'April', 'Mei', 'Jun', 
                   'Julai', 'Ogos', 'September', 'Oktober', 'November', 'Disember']
    p.setFont("Helvetica", 12)
    p.drawCentredString(width/2, height - 70, f"Bulan: {bulan_names[filter_month]} {filter_year}")

    total = len(records)
    lelaki = sum(1 for r in records if r.murid.jantina == 'Lelaki')
    perempuan = sum(1 for r in records if r.murid.jantina == 'Perempuan')

    y = height - 110
    p.setFont("Helvetica-Bold", 12)
    p.drawString(50, y, f"Jumlah Keseluruhan: {total}")
    p.drawString(250, y, f"Lelaki: {lelaki}")
    p.drawString(350, y, f"Perempuan: {perempuan}")

    y -= 40
    p.setFont("Helvetica-Bold", 10)
    headers = ['Tarikh', 'Masa', 'Nama', 'Kelas', 'Jantina']
    x_positions = [50, 120, 180, 380, 480]

    for i, header in enumerate(headers):
        p.drawString(x_positions[i], y, header)

    y -= 20
    p.setFont("Helvetica", 9)

    for record in records:
        if y < 50:
            p.showPage()
            y = height - 50
            p.setFont("Helvetica", 9)

        p.drawString(50, y, record.tarikh.strftime('%d/%m/%Y'))
        p.drawString(120, y, record.masa_sampai.strftime('%H:%M'))

        nama = record.murid.nama_penuh[:30]
        p.drawString(180, y, nama)
        p.drawString(380, y, record.murid.kelas.nama_kelas)
        p.drawString(480, y, record.murid.jantina)

        y -= 15

    p.save()
    buffer.seek(0)

    filename = f"Laporan_Kehadiran_Lewat_{filter_month}_{filter_year}.pdf"

    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/pdf'
    )

@app.route('/api/stats')
def api_stats():
    stats = get_weekly_stats()
    return jsonify(stats)

@app.errorhandler(404)
def not_found(e):
    return render_template('error.html', error='Halaman tidak dijumpai'), 404

@app.errorhandler(500)
def server_error(e):
    return render_template('error.html', error='Ralat pelayan'), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
